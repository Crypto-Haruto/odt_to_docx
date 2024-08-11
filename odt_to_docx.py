import sys
import os
from odf.opendocument import OpenDocument
from odf.text import P
from docx import Document

def convert_odt_to_docx(odt_file):
    try:
        # Load the ODT file
        odt_doc = OpenDocument(odt_file)
    except Exception as e:
        print(f"Error opening ODT file: {e}")
        return

    docx_file = os.path.splitext(odt_file)[0] + '.docx'
    docx_doc = Document()

    # Extract text from ODT and add to DOCX
    for paragraph in odt_doc.getElementsByType(P):
        if paragraph.firstChild is not None:
            docx_doc.add_paragraph(paragraph.firstChild.data)

    try:
        docx_doc.save(docx_file)
        print(f"Converted '{odt_file}' to '{docx_file}'.")
    except Exception as e:
        print(f"Error saving DOCX file: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python odt_to_docx_converter.py <path_to_odt_file>")
        sys.exit(1)

    input_file = sys.argv[1]

    if not os.path.isfile(input_file):
        print(f"File '{input_file}' does not exist.")
        sys.exit(1)

    convert_odt_to_docx(input_file)